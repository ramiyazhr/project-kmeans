import streamlit as st
from supabase import create_client
import pandas as pd
from sklearn.cluster import KMeans
from PIL import Image
import base64
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO

# === Inisialisasi session_state ===
if "data" not in st.session_state:
    st.session_state["data"] = pd.DataFrame(columns=["Wilayah", "Total Denda", "Bulan", "Tahun"])

# --- Koneksi Supabase ---
url = "https://xdxbqsiofkjjmhnypxku.supabase.co"
key = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InhkeGJxc2lvZmtqam1obnlweGt1Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3NTUxNzE5NjIsImV4cCI6MjA3MDc0Nzk2Mn0.dtQQLg76h2ZS7ZeDzfxkW00k3ix9Rs9BohkWU7Hhyxs"
supabase = create_client(url, key)

# --- Inisialisasi Session State ---
if "login_status" not in st.session_state:
    st.session_state.login_status = False
if "username" not in st.session_state:
    st.session_state.username = ""

# --- Fungsi Login ---
def login(username, password):
    res = supabase.table("users") \
        .select("*") \
        .eq("username", username) \
        .eq("password", password) \
        .execute()
    
    if res.data and len(res.data) > 0:
        st.session_state.login_status = True
        st.session_state.username = username
    else:
        st.error("Username atau password salah")

# --- Fungsi Logout ---
def logout():
    st.session_state.login_status = False
    st.session_state.username = ""

# --- Fungsi untuk set background ---
def set_bg(image_file):
    with open(image_file, "rb") as img:
        encoded = base64.b64encode(img.read()).decode()
    css = f"""
    <style>
    .stApp {{
        background: linear-gradient(rgba(255,255,255,0.6), rgba(255,255,255,0.6)),
                    url(data:image/jpeg;base64,{encoded});
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
    }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def menu_utama():
    # ====== Header sambutan dengan kotak background samar ======
    st.markdown(
        """
        <div style="
            background: rgba(255,255,255,0.7);
            padding: 16px 24px;
            border-radius: 12px;
            border: 1px solid rgba(0,0,0,0.06);
            box-shadow: 0 2px 10px rgba(0,0,0,0.04);
            margin-bottom: 16px;
        ">
            <h3 style="margin:0; text-align:center; color:#111; font-weight:400; line-height:1.4;">
                Selamat datang di Sistem <b>Implementasi Metode K-Means Clustering</b> Tunggakan Pajak Kendaraan Bermotor
            </h3>
        </div>
        """,
        unsafe_allow_html=True
    )

    # ====== Sidebar Menu ======
    st.sidebar.title("Menu")
    pilihan = st.sidebar.radio(
        "Pilih halaman",
        (
            "Input Data Tunggakan",
            "Klasterisasi K-Means",
            "Output Data"
        ),
        index=0
    )
    st.sidebar.button("Logout", on_click=logout)

    # ====== Konten Halaman (sementara placeholder dulu) ======
    if pilihan == "Input Data Tunggakan":
        st.subheader("💰 Input Data Tunggakan Pajak (Tahunan per Zona)")
    with st.form("form_tunggakan"):
        wilayah = st.text_input("Nama Wilayah")
        total_denda = st.number_input("Total Denda (Rp)", min_value=0, step=1000)
        tahun = st.number_input("Tahun", min_value=2020, max_value=2030, value=2024)
        submit = st.form_submit_button("Simpan Data")

    if submit:
        # cek apakah data untuk wilayah & tahun sudah ada
        existing = st.session_state.data[
            (st.session_state.data["Wilayah"] == wilayah) &
            (st.session_state.data["Tahun"] == tahun)
        ]
        if not existing.empty:
            st.error("⚠️ Data untuk wilayah & tahun ini sudah ada. Tidak bisa disimpan.")
        else:
            st.session_state.data = pd.concat(
                [st.session_state.data,
                 pd.DataFrame([[wilayah, total_denda, tahun]],
                              columns=["Wilayah", "Total Denda", "Tahun"])],
                ignore_index=True
            )
            st.success(f"✅ Data wilayah **{wilayah}** berhasil disimpan untuk tahun {tahun}!")

    elif pilihan == "Klasterisasi K-Means":
        st.subheader("📈 Klasterisasi Metode K-Means")

        if st.session_state["data"].empty:
            st.warning("⚠️ Belum ada data untuk diklasterisasi. Silakan isi dulu di menu Input Data.")
        else:
            st.write("### Data Tunggakan yang Tersimpan")
            st.dataframe(st.session_state["data"][["Wilayah", "Total Denda"]])

            k = st.slider("Pilih jumlah cluster (k)", min_value=2, max_value=6, value=3)

            if st.button("🚀 Proses Klasterisasi"):
                from sklearn.preprocessing import StandardScaler

                df = st.session_state["data"].copy()
                X = df[["Total Denda"]].values

                scaler = StandardScaler()
                X_scaled = scaler.fit_transform(X)

                kmeans = KMeans(n_clusters=k, random_state=42, n_init=10)
                labels = kmeans.fit_predict(X_scaled)

                df["Cluster"] = labels
                st.session_state["cluster_result"] = df

                st.success(f"✅ Klasterisasi selesai. Iterasi dilakukan sebanyak {kmeans.n_iter_} kali.")
                st.dataframe(df.sort_values("Cluster"))

                centroids_scaled = kmeans.cluster_centers_
                centroids_orig = scaler.inverse_transform(centroids_scaled).flatten()
                st.write("### 🎯 Centroid (Total Denda, skala asli)")
                st.table({f"Cluster {i}": [f"Rp {c:,.0f}"] for i, c in enumerate(centroids_orig)})

                st.info("➡️ Kamu bisa lanjut ke menu **Output Data** untuk melihat daftar zona per cluster.")

    elif pilihan == "Output Data":
        st.subheader("🧾 Laporan Hasil Klasterisasi Kendaraan Bermotor")
    st.markdown("#### Tahun 2024")

    # Cek apakah sudah ada hasil cluster
    if "cluster_result" not in st.session_state or st.session_state["cluster_result"].empty:
        st.warning("⚠️ Belum ada hasil klasterisasi. Silakan lakukan proses klasterisasi terlebih dahulu.")
    else:
        df_cluster = st.session_state["cluster_result"]

        # === Grafik Pie Chart ===
        st.markdown("### 📊 Grafik Distribusi Cluster")
        cluster_counts = df_cluster["Cluster"].value_counts()

        fig, ax = plt.subplots()
        ax.pie(cluster_counts, labels=cluster_counts.index, autopct="%1.1f%%", startangle=90)
        ax.axis("equal")  # biar bulat sempurna
        st.pyplot(fig)

        # === Tabel hasil cluster ===
        st.markdown("### 📑 Tabel Hasil Klasterisasi")
        st.dataframe(df_cluster[["Wilayah", "Total Denda", "Cluster"]])

        # Ringkasan jumlah cluster
        st.markdown("### 📊 Ringkasan Jumlah Zona per Cluster")
        cluster_summary = df_cluster.groupby("Cluster").size().reset_index(name="Jumlah Zona")
        st.table(cluster_summary)

        # Tombol download


        if st.button("⬇️ Unduh Laporan (Word)"):
            doc = Document()
            doc.add_heading("Laporan Hasil Klasterisasi Kendaraan Bermotor", 0)
            doc.add_paragraph("Tahun 2024\n")
            
            # Tambahkan ringkasan
            doc.add_heading("Ringkasan Klasterisasi", level=1)
            for i, row in cluster_summary.iterrows():
                doc.add_paragraph(f"Cluster {row['Cluster']}: {row['Jumlah Zona']} zona")
            
            # Tambahkan tabel
            doc.add_heading("Tabel Hasil Klasterisasi", level=1)
            table = doc.add_table(rows=1, cols=len(df_cluster.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df_cluster.columns):
                hdr_cells[i].text = col
            for _, row in df_cluster.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="⬇️ Download Word (.docx)",
                data=buffer,
                file_name="laporan_klasterisasi.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.success("✅ Hasil klasterisasi siap dicetak atau disimpan.")



# --- Halaman Login ---
if not st.session_state.login_status:
    # Set background samar
    set_bg("kantorsamsat.jpg")

    # Judul login (tebal & warna hitam)
    st.markdown(
        "<h2 style='text-align:center; color:black; font-weight:bold;'>Login Admin</h2>",
        unsafe_allow_html=True
    )

   # Label Username (bold & hitam)
    st.markdown("<label style='font-weight:bold; color:black;'>Username</label>", unsafe_allow_html=True)
    username = st.text_input("", key="username_input").strip()

    # Label Password (bold & hitam)
    st.markdown("<label style='font-weight:bold; color:black;'>Password</label>", unsafe_allow_html=True)
    password = st.text_input("", type="password", key="password_input").strip()

    if st.button("Login"):
        login(username, password)
else:

    menu_utama()




